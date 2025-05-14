"""
Optimized Document Processor Module

This module provides an optimized implementation for processing very large PDFs
and handling high-volume processing scenarios. It includes:

1. Parallel processing capabilities
2. Streaming extraction for large documents
3. Memory optimization techniques
4. Performance improvements for regex patterns
5. Distributed processing support
6. Queuing system for batch processing
"""

import os
import re
import json
import logging
import tempfile
import time
import threading
import queue
import multiprocessing as mp
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor
from typing import Dict, List, Any, Optional, Tuple, Union, Callable, Generator
import pandas as pd
import numpy as np
from datetime import datetime
import traceback
import psutil
import io

# Import enhanced processing modules
from .document_processor import DocumentProcessor

# Import libraries for document processing
try:
    import PyPDF2
    import fitz  # PyMuPDF
    from docx import Document as DocxDocument
    import tabula
    import camelot
    from PIL import Image
    import pytesseract
    HAS_ADVANCED_PROCESSING = True
except ImportError:
    HAS_ADVANCED_PROCESSING = False
    logging.warning("Advanced document processing libraries not available. Some features will be limited.")

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Define processing constants
DEFAULT_CHUNK_SIZE = 10  # Number of pages to process in a single chunk
DEFAULT_MAX_WORKERS = max(1, min(os.cpu_count() or 1, 8))  # Cap at 8 or number of CPUs
DEFAULT_MEMORY_LIMIT = 0.8  # Max memory usage (percentage of total)
DEFAULT_QUEUE_SIZE = 100   # Default size for processing queue

class PerformanceMetrics:
    """Track performance metrics for document processing."""
    
    def __init__(self):
        self.metrics = {
            "start_time": None,
            "end_time": None,
            "total_duration": None,
            "pages_processed": 0,
            "extraction_time": 0,
            "ocr_time": 0,
            "table_extraction_time": 0,
            "regex_time": 0, 
            "peak_memory_usage": 0,
            "stages": {}
        }
    
    def start(self):
        """Start timing."""
        self.metrics["start_time"] = time.time()
        
    def end(self):
        """End timing and calculate total duration."""
        self.metrics["end_time"] = time.time()
        self.metrics["total_duration"] = self.metrics["end_time"] - self.metrics["start_time"]
        
    def start_stage(self, stage_name: str):
        """Start timing a specific processing stage."""
        if stage_name not in self.metrics["stages"]:
            self.metrics["stages"][stage_name] = {"starts": [], "ends": [], "durations": []}
        
        self.metrics["stages"][stage_name]["starts"].append(time.time())
        
    def end_stage(self, stage_name: str):
        """End timing a specific processing stage."""
        if stage_name in self.metrics["stages"]:
            self.metrics["stages"][stage_name]["ends"].append(time.time())
            
            # Calculate duration
            start = self.metrics["stages"][stage_name]["starts"][-1]
            end = self.metrics["stages"][stage_name]["ends"][-1]
            duration = end - start
            self.metrics["stages"][stage_name]["durations"].append(duration)
            
    def record_memory_usage(self):
        """Record current memory usage."""
        process = psutil.Process(os.getpid())
        memory_info = process.memory_info()
        memory_usage = memory_info.rss / 1024 / 1024  # Convert to MB
        
        if memory_usage > self.metrics["peak_memory_usage"]:
            self.metrics["peak_memory_usage"] = memory_usage
            
    def increment_pages(self, count: int = 1):
        """Increment number of processed pages."""
        self.metrics["pages_processed"] += count
        
    def to_dict(self) -> Dict[str, Any]:
        """Convert metrics to dictionary for serialization."""
        # Calculate averages for stages
        for stage_name, stage_data in self.metrics["stages"].items():
            if stage_data["durations"]:
                stage_data["average_duration"] = sum(stage_data["durations"]) / len(stage_data["durations"])
                stage_data["total_duration"] = sum(stage_data["durations"])
            else:
                stage_data["average_duration"] = 0
                stage_data["total_duration"] = 0
                
        return self.metrics
    
    def summary(self) -> str:
        """Generate a human-readable summary of metrics."""
        if not self.metrics["end_time"]:
            self.end()
            
        summary = []
        summary.append(f"Total processing time: {self.metrics['total_duration']:.2f} seconds")
        summary.append(f"Pages processed: {self.metrics['pages_processed']}")
        summary.append(f"Pages per second: {self.metrics['pages_processed'] / max(1, self.metrics['total_duration']):.2f}")
        summary.append(f"Peak memory usage: {self.metrics['peak_memory_usage']:.2f} MB")
        
        # Add stage summaries
        summary.append("\nStage performance:")
        for stage_name, stage_data in self.metrics["stages"].items():
            if "total_duration" in stage_data:
                summary.append(f"  - {stage_name}: {stage_data['total_duration']:.2f}s total, "
                              f"{stage_data['average_duration']:.4f}s avg per call")
        
        return "\n".join(summary)


class PageStreamProcessor:
    """Process PDF pages in a streaming fashion to reduce memory usage."""
    
    def __init__(self, file_path: str):
        """Initialize with PDF file path."""
        self.file_path = file_path
        
    def stream_pages(self, chunk_size: int = DEFAULT_CHUNK_SIZE) -> Generator[List[Tuple[int, Any]], None, None]:
        """
        Stream PDF pages in chunks to reduce memory usage.
        
        Args:
            chunk_size: Number of pages to process in each chunk
            
        Yields:
            List of (page_num, page_object) tuples
        """
        try:
            doc = fitz.open(self.file_path)
            total_pages = len(doc)
            
            for i in range(0, total_pages, chunk_size):
                chunk_end = min(i + chunk_size, total_pages)
                pages = [(page_num, doc.load_page(page_num)) for page_num in range(i, chunk_end)]
                yield pages
                
                # Explicitly delete page objects after yielding to free memory
                for _, page in pages:
                    page = None
                    
                # Force garbage collection
                import gc
                gc.collect()
                
        except Exception as e:
            logger.error(f"Error streaming PDF pages: {e}", exc_info=True)
            
            # Fallback to PyPDF2 if PyMuPDF fails
            try:
                with open(self.file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    total_pages = len(reader.pages)
                    
                    for i in range(0, total_pages, chunk_size):
                        chunk_end = min(i + chunk_size, total_pages)
                        pages = [(page_num, reader.pages[page_num]) for page_num in range(i, chunk_end)]
                        yield pages
                        
                        # Clean up to free memory
                        for _, page in pages:
                            page = None
                            
                        # Force garbage collection
                        import gc
                        gc.collect()
                        
            except Exception as e2:
                logger.error(f"Fallback PDF streaming also failed: {e2}", exc_info=True)
                yield []
    
    def get_page_count(self) -> int:
        """Get the total number of pages in the PDF."""
        try:
            doc = fitz.open(self.file_path)
            count = len(doc)
            doc.close()
            return count
        except Exception as e:
            logger.error(f"Error getting page count with PyMuPDF: {e}", exc_info=True)
            
            # Fallback to PyPDF2 if PyMuPDF fails
            try:
                with open(self.file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    return len(reader.pages)
            except Exception as e2:
                logger.error(f"Fallback page count also failed: {e2}", exc_info=True)
                return 0


class ChunkedTextProcessor:
    """Process extracted text in chunks to optimize regex performance."""
    
    def __init__(self, patterns: Dict[str, re.Pattern]):
        """
        Initialize with regex patterns.
        
        Args:
            patterns: Dictionary of named regex patterns
        """
        self.patterns = patterns
        
    def process_text_in_chunks(self, text: str, chunk_size: int = 10000) -> Dict[str, List[str]]:
        """
        Process text in chunks for better regex performance.
        
        Args:
            text: Text to process
            chunk_size: Size of each text chunk
            
        Returns:
            Dictionary with pattern matches
        """
        results = {name: [] for name in self.patterns}
        
        # Process text in chunks
        for i in range(0, len(text), chunk_size):
            chunk = text[i:i + chunk_size]
            
            # Process each pattern
            for name, pattern in self.patterns.items():
                matches = pattern.findall(chunk)
                if matches:
                    results[name].extend(matches)
        
        # Remove duplicates
        for name in results:
            results[name] = list(set(results[name]))
            
        return results


class ParallelTableExtractor:
    """Extract tables from PDF in parallel."""
    
    def __init__(self, file_path: str, metrics: Optional[PerformanceMetrics] = None):
        """
        Initialize with PDF file path.
        
        Args:
            file_path: Path to the PDF file
            metrics: Performance metrics tracker
        """
        self.file_path = file_path
        self.metrics = metrics
        
    def extract_tables_parallel(self, max_workers: int = DEFAULT_MAX_WORKERS) -> List[Dict[str, Any]]:
        """
        Extract tables from PDF using parallel processing.
        
        Args:
            max_workers: Maximum number of worker processes
            
        Returns:
            List of extracted tables
        """
        if self.metrics:
            self.metrics.start_stage("table_extraction")
            
        all_tables = []
        
        try:
            # Get page count
            page_streamer = PageStreamProcessor(self.file_path)
            total_pages = page_streamer.get_page_count()
            
            # Configure process pool for CPU-bound task
            with ProcessPoolExecutor(max_workers=max_workers) as executor:
                # Process pages in batches for better performance
                batch_size = max(1, total_pages // max_workers)
                futures = []
                
                for i in range(0, total_pages, batch_size):
                    end_page = min(i + batch_size, total_pages)
                    page_range = f"{i+1}-{end_page}"
                    futures.append(executor.submit(self._extract_tables_from_pages, page_range))
                
                # Collect results
                for future in futures:
                    tables = future.result()
                    if tables:
                        all_tables.extend(tables)
        
        except Exception as e:
            logger.error(f"Error extracting tables in parallel: {e}", exc_info=True)
            
            # Fallback to sequential extraction
            logger.info("Falling back to sequential table extraction")
            try:
                tables = self._extract_tables_sequential()
                if tables:
                    all_tables.extend(tables)
            except Exception as e2:
                logger.error(f"Sequential table extraction also failed: {e2}", exc_info=True)
        
        if self.metrics:
            self.metrics.end_stage("table_extraction")
            
        return all_tables
    
    def _extract_tables_from_pages(self, pages: str) -> List[Dict[str, Any]]:
        """
        Extract tables from specific pages.
        
        Args:
            pages: Page range string (e.g., "1-5")
            
        Returns:
            List of extracted tables
        """
        tables = []
        
        try:
            # Use camelot for more accurate table detection
            camelot_tables = camelot.read_pdf(self.file_path, pages=pages, flavor='lattice')
            
            # Convert to common format
            for i, table in enumerate(camelot_tables):
                df = table.df
                
                tables.append({
                    "page": int(table.page) if hasattr(table, 'page') else 0,
                    "rows": len(df),
                    "columns": len(df.columns),
                    "headers": df.iloc[0].tolist() if not df.empty else [],
                    "data": df.iloc[1:].to_dict(orient='records') if len(df) > 1 else []
                })
                
        except Exception as e:
            logger.warning(f"Camelot failed for pages {pages}: {e}")
            
            # Fall back to tabula
            try:
                tabula_tables = tabula.read_pdf(self.file_path, pages=pages, multiple_tables=True)
                
                # Convert to common format
                for i, df in enumerate(tabula_tables):
                    if isinstance(df, pd.DataFrame) and not df.empty:
                        tables.append({
                            "page": int(pages.split('-')[0]) + i,  # Approximate page number
                            "rows": len(df),
                            "columns": len(df.columns),
                            "headers": df.columns.tolist(),
                            "data": df.to_dict(orient='records')
                        })
            except Exception as e2:
                logger.error(f"Tabula also failed for pages {pages}: {e2}")
        
        return tables
    
    def _extract_tables_sequential(self) -> List[Dict[str, Any]]:
        """
        Extract tables sequentially as a fallback.
        
        Returns:
            List of extracted tables
        """
        tables = []
        
        try:
            # Use tabula for table extraction
            tabula_tables = tabula.read_pdf(self.file_path, pages='all', multiple_tables=True)
            
            # Convert to common format
            for i, df in enumerate(tabula_tables):
                if isinstance(df, pd.DataFrame) and not df.empty:
                    tables.append({
                        "page": i + 1,  # Approximate page number
                        "rows": len(df),
                        "columns": len(df.columns),
                        "headers": df.columns.tolist(),
                        "data": df.to_dict(orient='records')
                    })
        except Exception as e:
            logger.error(f"Error extracting tables sequentially: {e}", exc_info=True)
        
        return tables


class MemoryOptimizedProcessor:
    """Memory-efficient document processor for very large PDFs."""
    
    def __init__(self, max_memory_percent: float = DEFAULT_MEMORY_LIMIT, metrics: Optional[PerformanceMetrics] = None):
        """
        Initialize with memory limits.
        
        Args:
            max_memory_percent: Maximum percentage of system memory to use
            metrics: Performance metrics tracker
        """
        self.max_memory_percent = max_memory_percent
        self.metrics = metrics
        
        # Define optimized regex patterns with precompilation
        self.patterns = {
            "isin": re.compile(r'\b[A-Z]{2}[A-Z0-9]{9}[0-9]\b'),
            "currency": re.compile(r'(?:USD|EUR|GBP|JPY|CHF|CAD|AUD|NZD|€|\$|£|¥)'),
            "amount": re.compile(r'(?:[\$€£¥])\s*\d+(?:,\d{3})*(?:\.\d+)?|\d+(?:,\d{3})*(?:\.\d+)?\s*(?:USD|EUR|GBP|JPY|CHF|CAD|AUD|NZD|million|billion|m|bn|k)'),
            "percentage": re.compile(r'\d+(?:\.\d+)?\s*%'),
            "date": re.compile(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2}|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}')
        }
        
        # Financial statement section patterns
        self.section_patterns = {
            "income_statement": re.compile(r'(?:income statement|statement of (?:income|earnings|operations|profit and loss)|profit and loss|p&l)', re.IGNORECASE),
            "balance_sheet": re.compile(r'(?:balance sheet|statement of (?:financial position|financial condition))', re.IGNORECASE),
            "cash_flow": re.compile(r'(?:cash flow|statement of cash flows)', re.IGNORECASE),
            "portfolio": re.compile(r'(?:portfolio|holdings|investments|securities|positions)', re.IGNORECASE)
        }
        
    def check_memory_usage(self) -> bool:
        """
        Check if memory usage is within limits.
        
        Returns:
            True if memory usage is OK, False if exceeded
        """
        if self.metrics:
            self.metrics.record_memory_usage()
            
        process = psutil.Process(os.getpid())
        memory_info = process.memory_info()
        memory_usage = memory_info.rss / 1024 / 1024  # Convert to MB
        
        total_memory = psutil.virtual_memory().total / 1024 / 1024  # Total memory in MB
        memory_limit = total_memory * self.max_memory_percent
        
        return memory_usage <= memory_limit
    
    def process_pdf_in_chunks(self, file_path: str, chunk_size: int = DEFAULT_CHUNK_SIZE) -> Dict[str, Any]:
        """
        Process large PDF in chunks to manage memory usage.
        
        Args:
            file_path: Path to the PDF file
            chunk_size: Number of pages to process in each chunk
            
        Returns:
            Extracted data from the PDF
        """
        if self.metrics:
            self.metrics.start_stage("chunked_pdf_processing")
            
        result = {
            "file_type": "pdf",
            "processed_at": datetime.now().isoformat(),
            "pages": 0,
            "text": "",
            "tables": [],
            "isins": [],
            "financial_data": {
                "currencies": [],
                "amounts": [],
                "percentages": [],
                "dates": [],
                "financial_terms": []
            },
            "metadata": {}
        }
        
        # Initialize page streamer
        page_streamer = PageStreamProcessor(file_path)
        result["pages"] = page_streamer.get_page_count()
        
        # Process metadata if possible
        try:
            doc = fitz.open(file_path)
            metadata = doc.metadata
            if metadata:
                result["metadata"] = {
                    "title": metadata.get("title", ""),
                    "author": metadata.get("author", ""),
                    "subject": metadata.get("subject", ""),
                    "keywords": metadata.get("keywords", ""),
                    "creator": metadata.get("creator", ""),
                    "producer": metadata.get("producer", ""),
                    "creationDate": metadata.get("creationDate", ""),
                    "modDate": metadata.get("modDate", "")
                }
            doc.close()
        except Exception as e:
            logger.warning(f"Failed to extract metadata: {e}")
            
        # Use StringIO to efficiently concatenate text
        text_buffer = io.StringIO()
        
        # Process each chunk of pages
        for pages in page_streamer.stream_pages(chunk_size):
            # Extract text from each page in the chunk
            for page_num, page in pages:
                if self.metrics:
                    self.metrics.start_stage("text_extraction")
                    
                try:
                    page_text = page.get_text()
                    text_buffer.write(page_text)
                    text_buffer.write("\n\n")
                    
                    if self.metrics:
                        self.metrics.increment_pages()
                        
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {e}")
                
                if self.metrics:
                    self.metrics.end_stage("text_extraction")
                    
            # Check memory usage and take action if needed
            if not self.check_memory_usage():
                logger.warning("Memory usage exceeded threshold. Processing partial data.")
                break
                
        # Extract tables in parallel
        table_extractor = ParallelTableExtractor(file_path, self.metrics)
        tables = table_extractor.extract_tables_parallel()
        result["tables"] = tables
        
        # Get the full text
        full_text = text_buffer.getvalue()
        
        # Process the text in chunks to extract patterns
        text_processor = ChunkedTextProcessor(self.patterns)
        
        if self.metrics:
            self.metrics.start_stage("pattern_extraction")
            
        pattern_results = text_processor.process_text_in_chunks(full_text)
        
        # Update result with pattern matches
        result["isins"] = pattern_results["isin"]
        result["financial_data"]["currencies"] = pattern_results["currency"]
        result["financial_data"]["amounts"] = pattern_results["amount"]
        result["financial_data"]["percentages"] = pattern_results["percentage"]
        result["financial_data"]["dates"] = pattern_results["date"]
        
        if self.metrics:
            self.metrics.end_stage("pattern_extraction")
            
        # Don't store the full text in the result to save memory
        # Instead, save to a temporary file if needed and reference it
        if len(full_text) > 1000000:  # If text is larger than ~1MB
            temp_dir = tempfile.gettempdir()
            text_file = os.path.join(temp_dir, f"pdf_text_{os.path.basename(file_path)}.txt")
            
            with open(text_file, 'w', encoding='utf-8', errors='ignore') as f:
                f.write(full_text)
                
            result["text_file"] = text_file
            result["text"] = full_text[:1000] + "... [truncated, see text_file]"
        else:
            result["text"] = full_text
            
        if self.metrics:
            self.metrics.end_stage("chunked_pdf_processing")
            
        return result


class TaskQueue:
    """Queue system for handling large batch processing jobs."""
    
    def __init__(self, max_size: int = DEFAULT_QUEUE_SIZE, num_workers: int = DEFAULT_MAX_WORKERS):
        """
        Initialize the task queue.
        
        Args:
            max_size: Maximum queue size
            num_workers: Number of worker threads
        """
        self.queue = queue.Queue(maxsize=max_size)
        self.results = {}
        self.num_workers = num_workers
        self.workers = []
        self.running = False
        self.lock = threading.Lock()
        
    def start_workers(self):
        """Start worker threads."""
        self.running = True
        
        for _ in range(self.num_workers):
            worker = threading.Thread(target=self._worker_loop)
            worker.daemon = True
            worker.start()
            self.workers.append(worker)
            
    def stop_workers(self):
        """Stop worker threads."""
        self.running = False
        
        # Add None tasks to unblock workers
        for _ in range(self.num_workers):
            self.queue.put(None)
            
        # Wait for all workers to finish
        for worker in self.workers:
            worker.join()
            
        self.workers = []
        
    def _worker_loop(self):
        """Worker thread main loop."""
        while self.running:
            task = self.queue.get()
            
            if task is None:
                self.queue.task_done()
                break
                
            try:
                task_id, task_func, args, kwargs = task
                result = task_func(*args, **kwargs)
                
                with self.lock:
                    self.results[task_id] = {
                        "status": "completed",
                        "result": result,
                        "error": None
                    }
            except Exception as e:
                error_traceback = traceback.format_exc()
                
                with self.lock:
                    self.results[task_id] = {
                        "status": "failed",
                        "result": None,
                        "error": str(e),
                        "traceback": error_traceback
                    }
            finally:
                self.queue.task_done()
                
    def add_task(self, task_id: str, task_func: Callable, *args, **kwargs) -> str:
        """
        Add a task to the queue.
        
        Args:
            task_id: Unique identifier for the task
            task_func: Function to execute
            *args: Positional arguments for the function
            **kwargs: Keyword arguments for the function
            
        Returns:
            Task identifier
        """
        with self.lock:
            self.results[task_id] = {
                "status": "queued",
                "result": None,
                "error": None
            }
            
        self.queue.put((task_id, task_func, args, kwargs))
        return task_id
        
    def get_result(self, task_id: str) -> Dict[str, Any]:
        """
        Get the result of a task.
        
        Args:
            task_id: Task identifier
            
        Returns:
            Task result information
        """
        with self.lock:
            return self.results.get(task_id, {"status": "unknown", "result": None, "error": "Task not found"})
            
    def wait_for_all(self):
        """Wait for all tasks to complete."""
        self.queue.join()
    
    def get_queue_status(self) -> Dict[str, Any]:
        """
        Get status of the queue.
        
        Returns:
            Queue status information
        """
        with self.lock:
            pending = sum(1 for info in self.results.values() if info["status"] == "queued")
            completed = sum(1 for info in self.results.values() if info["status"] == "completed")
            failed = sum(1 for info in self.results.values() if info["status"] == "failed")
            
            return {
                "queue_size": self.queue.qsize(),
                "pending_tasks": pending,
                "completed_tasks": completed,
                "failed_tasks": failed,
                "total_tasks": len(self.results)
            }


class OptimizedDocumentProcessor:
    """
    Enhanced document processor optimized for large PDFs and high-volume processing.
    Includes parallel processing, streaming extraction, and memory optimization.
    """
    
    def __init__(self, upload_dir: str = "uploads", max_workers: int = DEFAULT_MAX_WORKERS, 
                 memory_limit: float = DEFAULT_MEMORY_LIMIT):
        """
        Initialize the optimized document processor.
        
        Args:
            upload_dir: Directory for uploaded files
            max_workers: Maximum number of worker processes/threads
            memory_limit: Maximum memory usage as percentage of total
        """
        self.upload_dir = upload_dir
        os.makedirs(upload_dir, exist_ok=True)
        
        self.max_workers = max_workers
        self.memory_limit = memory_limit
        
        # Create base processor for non-critical operations
        self.base_processor = DocumentProcessor(upload_dir)
        
        # Initialize task queue for batch processing
        self.task_queue = TaskQueue(num_workers=max_workers)
        self.task_queue.start_workers()
        
        # Initialize metrics tracker
        self.metrics = PerformanceMetrics()
        
        logger.info(f"Initialized OptimizedDocumentProcessor with {max_workers} workers "
                   f"and {memory_limit:.0%} memory limit")
        
    def __del__(self):
        """Clean up resources on destruction."""
        if hasattr(self, 'task_queue'):
            self.task_queue.stop_workers()
        
    def process_document(self, file_path: str, file_type: str, 
                        processing_options: Optional[Dict] = None) -> Dict[str, Any]:
        """
        Process a document with optimized handling for large files.
        
        Args:
            file_path: Path to the document
            file_type: Type of the document
            processing_options: Options for processing
            
        Returns:
            Extracted document data
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        # Default processing options
        if processing_options is None:
            processing_options = {}
            
        # Set performance tracking
        self.metrics = PerformanceMetrics()
        self.metrics.start()
        
        # Get file extension
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        # Choose processing method based on file type and size
        file_size = os.path.getsize(file_path) / (1024 * 1024)  # Convert to MB
        
        result = None
        
        try:
            # Log processing start
            logger.info(f"Processing {file_type} document {file_path} ({file_size:.2f} MB)")
            
            # Process based on file type with optimized methods for large files
            if ext in ['.xlsx', '.xls'] and file_size > 50:
                result = self._process_large_excel(file_path, processing_options)
            elif ext == '.csv' and file_size > 50:
                result = self._process_large_csv(file_path, processing_options)
            elif ext == '.pdf':
                # Use optimized PDF processing for any PDF over 10MB or with explicitly set "large_file" option
                if file_size > 10 or processing_options.get('large_file', False):
                    result = self._process_large_pdf(file_path, processing_options)
                else:
                    # For smaller PDFs, add basic performance tracking to standard processor
                    self.metrics.start_stage("standard_pdf_processing")
                    result = self.base_processor._process_pdf(file_path, processing_options)
                    self.metrics.end_stage("standard_pdf_processing")
            elif ext in ['.doc', '.docx'] and file_size > 20:
                result = self._process_large_word(file_path, processing_options)
            else:
                # For smaller files, use standard processor with added metrics
                if ext in ['.xlsx', '.xls']:
                    self.metrics.start_stage("standard_excel_processing")
                    result = self.base_processor._process_excel(file_path, processing_options)
                    self.metrics.end_stage("standard_excel_processing")
                elif ext == '.csv':
                    self.metrics.start_stage("standard_csv_processing")
                    result = self.base_processor._process_csv(file_path, processing_options)
                    self.metrics.end_stage("standard_csv_processing")
                elif ext in ['.doc', '.docx']:
                    self.metrics.start_stage("standard_word_processing")
                    result = self.base_processor._process_word(file_path, processing_options)
                    self.metrics.end_stage("standard_word_processing")
                else:
                    self.metrics.start_stage("standard_text_processing")
                    result = self.base_processor._process_text(file_path, processing_options)
                    self.metrics.end_stage("standard_text_processing")
                    
            # Add performance metrics to result
            self.metrics.end()
            result["performance_metrics"] = self.metrics.to_dict()
            
            # Log processing summary
            logger.info(f"Completed processing {file_path}")
            logger.info(self.metrics.summary())
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing document: {e}", exc_info=True)
            
            # End metrics if still running
            if hasattr(self.metrics, 'end'):
                self.metrics.end()
                
            return {
                "file_type": file_type,
                "processed_at": datetime.now().isoformat(),
                "error": str(e),
                "performance_metrics": self.metrics.to_dict() if hasattr(self.metrics, 'to_dict') else {}
            }
            
    def process_document_async(self, file_path: str, file_type: str, 
                             processing_options: Optional[Dict] = None) -> str:
        """
        Process a document asynchronously.
        
        Args:
            file_path: Path to the document
            file_type: Type of the document
            processing_options: Options for processing
            
        Returns:
            Task ID for retrieving results
        """
        task_id = f"task_{int(time.time())}_{os.path.basename(file_path)}"
        self.task_queue.add_task(task_id, self.process_document, file_path, file_type, processing_options)
        return task_id
        
    def get_task_result(self, task_id: str) -> Dict[str, Any]:
        """
        Get result of an asynchronous task.
        
        Args:
            task_id: Task identifier
            
        Returns:
            Task result or status information
        """
        return self.task_queue.get_result(task_id)
        
    def get_queue_status(self) -> Dict[str, Any]:
        """
        Get status of the processing queue.
        
        Returns:
            Queue status information
        """
        return self.task_queue.get_queue_status()
        
    def process_batch(self, file_paths: List[str], file_types: List[str], 
                     processing_options: Optional[List[Dict]] = None) -> List[str]:
        """
        Process a batch of documents asynchronously.
        
        Args:
            file_paths: List of file paths
            file_types: List of file types
            processing_options: List of processing options
            
        Returns:
            List of task IDs
        """
        if len(file_paths) != len(file_types):
            raise ValueError("Number of file paths must match number of file types")
            
        if processing_options is None:
            processing_options = [None] * len(file_paths)
        elif len(processing_options) != len(file_paths):
            raise ValueError("Number of processing options must match number of file paths")
            
        task_ids = []
        
        for i, (file_path, file_type) in enumerate(zip(file_paths, file_types)):
            options = processing_options[i]
            task_id = self.process_document_async(file_path, file_type, options)
            task_ids.append(task_id)
            
        return task_ids
        
    def wait_for_batch(self, task_ids: List[str], timeout: Optional[float] = None) -> Dict[str, Dict[str, Any]]:
        """
        Wait for a batch of tasks to complete.
        
        Args:
            task_ids: List of task IDs
            timeout: Maximum wait time in seconds
            
        Returns:
            Dictionary of task results
        """
        start_time = time.time()
        results = {}
        
        # Loop until all tasks are done or timeout
        while len(results) < len(task_ids):
            # Check if timeout has been reached
            if timeout is not None and time.time() - start_time > timeout:
                break
                
            # Check status of each pending task
            for task_id in task_ids:
                if task_id not in results:
                    status = self.get_task_result(task_id)
                    if status["status"] in ["completed", "failed"]:
                        results[task_id] = status
                        
            # Sleep briefly to avoid tight loop
            if len(results) < len(task_ids):
                time.sleep(0.1)
                
        # Get any remaining results
        for task_id in task_ids:
            if task_id not in results:
                results[task_id] = self.get_task_result(task_id)
                
        return results
                
    def _process_large_pdf(self, file_path: str, options: Dict) -> Dict[str, Any]:
        """
        Process large PDF files with memory optimization and parallel processing.
        
        Args:
            file_path: Path to the PDF file
            options: Processing options
            
        Returns:
            Extracted PDF data
        """
        logger.info(f"Processing large PDF: {file_path}")
        
        # Create memory optimized processor
        memory_processor = MemoryOptimizedProcessor(self.memory_limit, self.metrics)
        
        # Process the PDF in chunks
        chunk_size = options.get('chunk_size', DEFAULT_CHUNK_SIZE)
        result = memory_processor.process_pdf_in_chunks(file_path, chunk_size)
        
        # Extract financial data if requested
        if options.get('extract_financial_data', True) and result.get("text"):
            self.metrics.start_stage("financial_data_extraction")
            
            # Use the base processor for financial data extraction (no need to optimize this step)
            financial_data = self.base_processor._extract_financial_data(result["text"])
            result["financial_data"].update(financial_data)
            
            self.metrics.end_stage("financial_data_extraction")
            
        return result
        
    def _process_large_excel(self, file_path: str, options: Dict) -> Dict[str, Any]:
        """
        Process large Excel files with memory optimization.
        
        Args:
            file_path: Path to the Excel file
            options: Processing options
            
        Returns:
            Extracted Excel data
        """
        logger.info(f"Processing large Excel file: {file_path}")
        self.metrics.start_stage("large_excel_processing")
        
        result = {
            "file_type": "excel",
            "processed_at": datetime.now().isoformat(),
            "sheets": [],
            "tables": [],
            "isins": []
        }
        
        try:
            # Use chunks to read Excel file
            excel_file = pd.ExcelFile(file_path)
            
            # Determine which sheets to process
            sheet_names = options.get('sheet_names', 'all')
            if sheet_names == 'all':
                sheet_names = excel_file.sheet_names
            elif isinstance(sheet_names, str):
                sheet_names = [sheet_names]
                
            # Process each sheet with chunking for large files
            for sheet_name in sheet_names:
                try:
                    self.metrics.start_stage(f"excel_sheet_{sheet_name}")
                    
                    # Get sheet dimensions first
                    df_info = pd.read_excel(file_path, sheet_name=sheet_name, nrows=0)
                    
                    # Extract basic sheet info
                    sheet_info = {
                        "name": sheet_name,
                        "columns": len(df_info.columns),
                        "column_names": df_info.columns.tolist()
                    }
                    
                    # Process in chunks to limit memory usage
                    chunk_size = 1000
                    chunks = []
                    row_count = 0
                    
                    for chunk in pd.read_excel(file_path, sheet_name=sheet_name, chunksize=chunk_size):
                        row_count += len(chunk)
                        
                        # Extract ISINs from the chunk
                        isins = self.base_processor._extract_isins_from_data(chunk)
                        if isins:
                            result["isins"].extend(isins)
                            
                        # Keep only first chunk for table preview
                        if not chunks:
                            chunks.append(chunk)
                            
                    # Update sheet info with row count
                    sheet_info["rows"] = row_count
                    result["sheets"].append(sheet_info)
                    
                    # Extract tables if requested (using only the first chunk for preview)
                    if options.get('extract_tables', True) and chunks:
                        table_info = {
                            "sheet": sheet_name,
                            "rows": row_count,
                            "columns": len(df_info.columns),
                            "headers": df_info.columns.tolist(),
                            "data": chunks[0].head(10).to_dict(orient='records')  # First 10 rows for preview
                        }
                        result["tables"].append(table_info)
                        
                    self.metrics.end_stage(f"excel_sheet_{sheet_name}")
                    
                except Exception as e:
                    logger.error(f"Error processing sheet {sheet_name}: {e}")
                    
            # Remove duplicates from ISINs
            result["isins"] = list(set(result["isins"]))
            
            self.metrics.end_stage("large_excel_processing")
            return result
            
        except Exception as e:
            logger.error(f"Error processing large Excel file: {e}")
            self.metrics.end_stage("large_excel_processing")
            return {
                "file_type": "excel",
                "processed_at": datetime.now().isoformat(),
                "error": str(e)
            }
            
    def _process_large_csv(self, file_path: str, options: Dict) -> Dict[str, Any]:
        """
        Process large CSV files with memory optimization.
        
        Args:
            file_path: Path to the CSV file
            options: Processing options
            
        Returns:
            Extracted CSV data
        """
        logger.info(f"Processing large CSV file: {file_path}")
        self.metrics.start_stage("large_csv_processing")
        
        result = {
            "file_type": "csv",
            "processed_at": datetime.now().isoformat(),
            "tables": [],
            "isins": []
        }
        
        try:
            # Get CSV info first
            df_info = pd.read_csv(file_path, nrows=0)
            result["columns"] = len(df_info.columns)
            result["column_names"] = df_info.columns.tolist()
            
            # Process in chunks to limit memory usage
            chunk_size = 10000
            chunks = []
            row_count = 0
            
            for chunk in pd.read_csv(file_path, chunksize=chunk_size):
                row_count += len(chunk)
                
                # Extract ISINs from the chunk
                isins = self.base_processor._extract_isins_from_data(chunk)
                if isins:
                    result["isins"].extend(isins)
                    
                # Keep only first chunk for table preview
                if not chunks:
                    chunks.append(chunk)
                    
            # Update result with row count
            result["rows"] = row_count
            
            # Extract tables if requested (using only the first chunk for preview)
            if options.get('extract_tables', True) and chunks:
                table_info = {
                    "rows": row_count,
                    "columns": len(df_info.columns),
                    "headers": df_info.columns.tolist(),
                    "data": chunks[0].head(10).to_dict(orient='records')  # First 10 rows for preview
                }
                result["tables"].append(table_info)
                
            # Remove duplicates from ISINs
            result["isins"] = list(set(result["isins"]))
            
            self.metrics.end_stage("large_csv_processing")
            return result
            
        except Exception as e:
            logger.error(f"Error processing large CSV file: {e}")
            self.metrics.end_stage("large_csv_processing")
            return {
                "file_type": "csv",
                "processed_at": datetime.now().isoformat(),
                "error": str(e)
            }
            
    def _process_large_word(self, file_path: str, options: Dict) -> Dict[str, Any]:
        """
        Process large Word documents with memory optimization.
        
        Args:
            file_path: Path to the Word document
            options: Processing options
            
        Returns:
            Extracted Word document data
        """
        logger.info(f"Processing large Word document: {file_path}")
        self.metrics.start_stage("large_word_processing")
        
        result = {
            "file_type": "word",
            "processed_at": datetime.now().isoformat(),
            "text": "",
            "tables": [],
            "isins": [],
            "metadata": {}
        }
        
        try:
            if not HAS_ADVANCED_PROCESSING:
                raise ImportError("Advanced processing libraries required for large Word documents")
                
            # Use python-docx to extract content
            doc = DocxDocument(file_path)
            
            # Extract text in chunks to limit memory usage
            text_buffer = io.StringIO()
            
            # Process paragraphs in batches
            batch_size = 100
            para_batches = [doc.paragraphs[i:i+batch_size] for i in range(0, len(doc.paragraphs), batch_size)]
            
            for batch in para_batches:
                for para in batch:
                    if para.text.strip():
                        text_buffer.write(para.text)
                        text_buffer.write("\n")
                        
                # Check memory usage
                if not MemoryOptimizedProcessor(self.memory_limit).check_memory_usage():
                    logger.warning("Memory usage exceeded threshold during Word processing")
                    break
                    
            # Get the full text
            full_text = text_buffer.getvalue()
            
            # Don't store the full text in the result to save memory if it's too large
            if len(full_text) > 1000000:  # If text is larger than ~1MB
                temp_dir = tempfile.gettempdir()
                text_file = os.path.join(temp_dir, f"word_text_{os.path.basename(file_path)}.txt")
                
                with open(text_file, 'w', encoding='utf-8', errors='ignore') as f:
                    f.write(full_text)
                    
                result["text_file"] = text_file
                result["text"] = full_text[:1000] + "... [truncated, see text_file]"
            else:
                result["text"] = full_text
                
            # Extract tables if present, with memory management
            if doc.tables:
                for i, table in enumerate(doc.tables):
                    # Skip if we've processed too many tables
                    if i >= 20:  # Arbitrary limit for large docs
                        break
                        
                    table_data = []
                    headers = []
                    
                    # Extract headers from first row
                    if table.rows:
                        headers = [cell.text for cell in table.rows[0].cells]
                        
                    # Extract data from remaining rows
                    max_rows = min(len(table.rows), 100)  # Limit to 100 rows for large tables
                    
                    for row_idx in range(1, max_rows):
                        row = table.rows[row_idx]
                        
                        row_data = {}
                        for col_idx, cell in enumerate(row.cells):
                            header = headers[col_idx] if col_idx < len(headers) else f"Column {col_idx+1}"
                            row_data[header] = cell.text
                            
                        table_data.append(row_data)
                        
                    result["tables"].append({
                        "id": f"table_{i+1}",
                        "headers": headers,
                        "rows": len(table.rows) - 1 if headers else len(table.rows),
                        "columns": len(headers),
                        "data": table_data
                    })
                    
                    # Check memory after each table
                    if not MemoryOptimizedProcessor(self.memory_limit).check_memory_usage():
                        logger.warning("Memory usage exceeded threshold during Word table processing")
                        break
                        
            # Extract metadata if available
            if hasattr(doc, 'core_properties') and doc.core_properties:
                result["metadata"] = {
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "subject": doc.core_properties.subject or "",
                    "keywords": doc.core_properties.keywords or "",
                    "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                    "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
                    "last_modified_by": doc.core_properties.last_modified_by or ""
                }
                
            # Extract ISINs from the text - use chunked processor for efficiency
            text_processor = ChunkedTextProcessor({"isin": re.compile(r'\b[A-Z]{2}[A-Z0-9]{9}[0-9]\b')})
            pattern_results = text_processor.process_text_in_chunks(full_text)
            result["isins"] = pattern_results["isin"]
            
            # Extract financial data with optimized pattern matching
            self.metrics.start_stage("financial_pattern_extraction")
            optimized_processor = MemoryOptimizedProcessor(self.memory_limit, self.metrics)
            text_processor = ChunkedTextProcessor(optimized_processor.patterns)
            pattern_results = text_processor.process_text_in_chunks(full_text)
            
            result["financial_data"] = {
                "currencies": pattern_results["currency"],
                "amounts": pattern_results["amount"],
                "percentages": pattern_results["percentage"],
                "dates": pattern_results["date"]
            }
            self.metrics.end_stage("financial_pattern_extraction")
            
            self.metrics.end_stage("large_word_processing")
            return result
            
        except Exception as e:
            logger.error(f"Error processing large Word document: {e}", exc_info=True)
            self.metrics.end_stage("large_word_processing")
            return {
                "file_type": "word",
                "processed_at": datetime.now().isoformat(),
                "error": str(e)
            }

# Helper function to measure improvement
def compare_performance(original_processor, optimized_processor, file_path, file_type, options=None):
    """
    Compare performance between original and optimized processors.
    
    Args:
        original_processor: Original document processor
        optimized_processor: Optimized document processor
        file_path: Path to test file
        file_type: Type of file
        options: Processing options
        
    Returns:
        Dictionary with performance comparison
    """
    # Process with original processor
    start_time_original = time.time()
    result_original = original_processor.process_document(file_path, file_type, options)
    end_time_original = time.time()
    original_duration = end_time_original - start_time_original
    
    # Process with optimized processor
    start_time_optimized = time.time()
    result_optimized = optimized_processor.process_document(file_path, file_type, options)
    end_time_optimized = time.time()
    optimized_duration = end_time_optimized - start_time_optimized
    
    # Calculate improvement
    speedup = original_duration / max(0.001, optimized_duration)
    percent_improvement = ((original_duration - optimized_duration) / original_duration) * 100
    
    # Get memory usage from optimized result if available
    memory_usage = None
    if "performance_metrics" in result_optimized:
        memory_usage = result_optimized["performance_metrics"].get("peak_memory_usage")
    
    comparison = {
        "file_path": file_path,
        "file_type": file_type,
        "original_duration": original_duration,
        "optimized_duration": optimized_duration,
        "speedup_factor": speedup,
        "percent_improvement": percent_improvement,
        "memory_usage_mb": memory_usage,
        "optimization_details": result_optimized.get("performance_metrics", {}).get("stages", {})
    }
    
    return comparison

# Example usage
if __name__ == "__main__":
    # This will only run when the script is executed directly
    import sys
    import os
    
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        file_type = sys.argv[2] if len(sys.argv) > 2 else "pdf"
        
        # Create processors
        from financial_document_processor import FinancialDocumentProcessor
        
        original_processor = FinancialDocumentProcessor()
        optimized_processor = OptimizedDocumentProcessor()
        
        # Run performance comparison
        comparison = compare_performance(original_processor, optimized_processor, file_path, file_type)
        
        print(f"Performance comparison results:")
        print(f"  Original processing time: {comparison['original_duration']:.2f} seconds")
        print(f"  Optimized processing time: {comparison['optimized_duration']:.2f} seconds")
        print(f"  Speedup factor: {comparison['speedup_factor']:.2f}x")
        print(f"  Improvement: {comparison['percent_improvement']:.2f}%")
        if comparison['memory_usage_mb']:
            print(f"  Peak memory usage: {comparison['memory_usage_mb']:.2f} MB")
        
        # Save detailed results
        with open("optimization_comparison.json", "w") as f:
            json.dump(comparison, f, indent=2)
            
        print(f"Detailed results saved to optimization_comparison.json")
    else:
        print("Please provide a file path to process")